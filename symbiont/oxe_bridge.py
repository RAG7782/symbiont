"""
OXÉ Bridge — SYMBIONT ↔ OXÉ API Integration

Conecta o organismo SYMBIONT ao OXÉ API, mapeando as castes do SYMBIONT
para os agentes especializados do OXÉ.

Mapeamento Caste → Agente OXÉ:
  Scout (Tier 0)  → busca paralela: jurisprudência + legislação + doutrina
  Media (Tier 1)  → construção: LexiaForge (redação) + JurisGuard (revisão)
  Major (Tier 2)  → estratégia: OXÉ Curador (jurimetria) + Neuropersuasão
  Minima (Suport) → formatação: Design Jurídico (.docx/.pdf)

Pipeline Premium (3 waves):
  Wave 1 (Scout)  → pesquisa paralela em todos os corpora
  Wave 2 (Media+Major) → análise estratégica + redação + blindagem
  Wave 3 (Minima) → formatação ABNT + entrega .docx

Configuração via env vars (nenhuma credencial hardcoded):
  OXE_URL       — URL base da API  (default: http://localhost:8500)
  OXE_EMAIL     — e-mail de login
  OXE_PASSWORD  — senha de login

Uso:
    from symbiont.oxe_bridge import OXEBridge

    bridge = OXEBridge()                        # lê env vars
    result = await bridge.run_premium(
        "elabore recurso sobre ISS na prestação de serviços"
    )
    print(result.wave3_peca)
"""

from __future__ import annotations

import asyncio
import logging
import os
import tempfile
import time
from dataclasses import dataclass, field


# ─── FastAPI request model (módulo-level para Pydantic v2) ───────────────────
try:
    from pydantic import BaseModel as _PydanticBase
    class PremiumRequest(_PydanticBase):
        query:                str
        conversation_id:      str  = ''
        estilo_advogado:      str  = ''
        use_assembly:         bool = False  # True → Wave 2 via LegalAssembly paralelo
        adversarial_persona:  str  = 'Procurador da Fazenda Nacional'
except ImportError:
    PremiumRequest = None  # type: ignore
from typing import Any

import requests

logger = logging.getLogger(__name__)

# ─── Configuração via env vars ─────────────────────────────────────────────────

OXE_DEFAULT_URL = "http://localhost:8500"
TOKEN_TTL = 82800  # 23h — renova antes de expirar (token válido por 24h)


# ─── Modelo de custo por tarefa ───────────────────────────────────────────────

COST_TIERS: dict[str, dict] = {
    "free":    {"model": "ollama/qwen3:1.7b",   "cost_usd": 0.000},
    "cheap":   {"model": "glm-5",              "cost_usd": 0.010},
    "medium":  {"model": "glm-5-turbo",        "cost_usd": 0.012},
    "high":    {"model": "glm-5.1",            "cost_usd": 0.015},
    "premium": {"model": "claude-sonnet-4-6",  "cost_usd": 0.060},
}

WAVE_COST: dict[str, str] = {
    "wave1_search":     "free",    # busca → Ollama local grátis
    "wave2_strategy":   "medium",  # jurimetria → GLM médio (glm-5 retorna vazio)
    "wave2_persuasion": "medium",  # neuropersuasão → GLM médio
    "wave2_draft":      "high",    # redação → GLM alto ou Claude
    "wave2_review":     "medium",  # revisão → GLM médio (glm-5 retorna vazio)
    "wave3_format":     "medium",  # formatação → glm-5-turbo (glm-5 retorna vazio)
}


# ─── Resultado do pipeline ────────────────────────────────────────────────────

@dataclass
class PremiumResult:
    query:                str
    wave1_jurisprudencia: list[dict] = field(default_factory=list)
    wave1_legislacao:     list[dict] = field(default_factory=list)
    wave1_doutrina:       list[dict] = field(default_factory=list)
    wave2_jurimetria:     str = ""
    wave2_persuasao:      str = ""
    wave2_rascunho:       str = ""
    wave2_revisado:       str = ""
    # Assembly mode: campos adicionais (preenchidos quando use_assembly=True)
    wave2_assembly_sintese:    str   = ""
    wave2_assembly_adversarial:str   = ""
    wave2_assembly_quality:    float = 0.0
    wave2_assembly_models:     dict  = field(default_factory=dict)
    wave2_mode:                str   = "sequential"  # "sequential" | "assembly"
    wave3_peca:           str = ""
    docx_path:            str = ""
    custo_estimado_usd:   float = 0.0
    tempo_total_s:        float = 0.0
    erros:                list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {k: v for k, v in self.__dict__.items()}


# ─── OXÉ Bridge ──────────────────────────────────────────────────────────────

class OXEBridge:
    """
    Bridge entre SYMBIONT e OXÉ API.

    Gerencia autenticação JWT com renovação automática (TOKEN_TTL = 23h).
    Todas as credenciais vêm de env vars — nenhuma hardcoded.
    """

    def __init__(
        self,
        oxe_url: str | None = None,
        email: str | None = None,
        password: str | None = None,
    ) -> None:
        self.oxe_url   = (oxe_url or os.environ.get("OXE_URL", OXE_DEFAULT_URL)).rstrip("/")
        self._email    = email    or os.environ.get("OXE_EMAIL", "")
        self._password = password or os.environ.get("OXE_PASSWORD", "")
        self._token    = ""
        self._token_ts = 0.0
        self._session  = requests.Session()
        self._session.headers.update({"Content-Type": "application/json"})

    # ── Auth ──────────────────────────────────────────────────────────────────

    def _get_token(self) -> str:
        """Retorna token JWT válido, renovando se expirado."""
        if self._token and (time.time() - self._token_ts) < TOKEN_TTL:
            return self._token
        if not self._email or not self._password:
            raise RuntimeError(
                "OXE_EMAIL e OXE_PASSWORD devem estar definidos como env vars"
            )
        r = self._session.post(
            f"{self.oxe_url}/login",
            json={"email": self._email, "password": self._password},
            timeout=15,
        )
        r.raise_for_status()
        self._token    = r.json()["token"]
        self._token_ts = time.time()
        logger.debug("oxe_bridge: token renovado")
        return self._token

    def _headers(self) -> dict[str, str]:
        return {"Authorization": f"Bearer {self._get_token()}"}

    # ── Chamadas OXÉ API ──────────────────────────────────────────────────────

    def buscar(
        self, query: str, collection: str = "oxe_jurisprudencia", top_k: int = 5
    ) -> list[dict]:
        """Busca semantica direta no Qdrant — sem loop pelo gateway."""
        import os as _os
        qdrant_url = _os.environ.get("QDRANT_URL", "http://qdrant:6333")
        ollama_url = _os.environ.get("OLLAMA_URL", "http://host.docker.internal:11434")
        try:
            er = self._session.post(
                f"{ollama_url}/api/embeddings",
                json={"model": "all-minilm", "prompt": query[:400]},
                timeout=15,
            )
            er.raise_for_status()
            vector = er.json()["embedding"]
            qr = self._session.post(
                f"{qdrant_url}/collections/{collection}/points/query",
                json={"query": vector, "limit": top_k, "with_payload": True},
                timeout=20,
            )
            qr.raise_for_status()
            points = qr.json().get("result", {}).get("points", [])
            results = []
            for p in points:
                pl = p.get("payload", {})
                results.append({
                    "tribunal": pl.get("sigla") or pl.get("tribunal") or pl.get("uf", "?"),
                    "data": pl.get("data", "?"),
                    "ementa": pl.get("ementa", ""),
                    "numero": pl.get("numero", ""),
                    "assunto": pl.get("assunto", ""),
                })
            logger.info("oxe_bridge: busca direta %s -> %d", collection, len(results))
            return results
        except Exception as exc:
            logger.warning("oxe_bridge: busca falhou %s: %s", collection, exc)
            return []

    def chat(
        self,
        message: str,
        skill: str | None = None,
        conversation_id: str = "premium",
        tier: str = "high",
    ) -> str:
        """Chama Z.ai/GLM com retry (429) + fallback glm-5 + Groq (2 modelos) + OpenRouter."""
        import os as _os
        import time as _time

        zai_key = _os.environ.get("ZAI_API_KEY", "")
        model_map = {
            "free": "glm-5", "cheap": "glm-5",
            "medium": "glm-5-turbo", "high": "glm-5.1", "premium": "glm-5.1",
        }
        model = model_map.get(tier, "glm-5.1")

        # Retry com backoff exponencial em 429
        _backoff = [10, 20, 40]
        for attempt, wait in enumerate([0] + _backoff):
            if wait:
                _time.sleep(wait)
            try:
                r = self._session.post(
                    "https://api.z.ai/api/paas/v4/chat/completions",
                    headers={"Authorization": f"Bearer {zai_key}", "Content-Type": "application/json"},
                    json={
                        "model": model,
                        "messages": [{"role": "user", "content": message}],
                        "max_tokens": 2000,
                        "temperature": 0.3,
                    },
                    timeout=90,
                )
                if r.status_code == 429:
                    retry_after = int(r.headers.get("Retry-After", wait or 3))
                    logger.warning("oxe_bridge: 429 Z.ai tier=%s attempt=%d wait=%ds", tier, attempt, retry_after)
                    if attempt < len(_backoff):
                        _time.sleep(retry_after)
                        continue
                    break  # esgotou retries — vai para fallback
                r.raise_for_status()
                content = r.json()["choices"][0]["message"]["content"]
                logger.info("oxe_bridge: GLM/%s -> %d chars", model, len(content))
                return content
            except Exception as exc:
                if "429" in str(exc) and attempt < len(_backoff):
                    continue
                logger.warning("oxe_bridge: GLM falhou tier=%s attempt=%d: %s", tier, attempt, exc)
                break

        # Fallback 1: Z.ai glm-5 (modelo mais leve, rate limit separado de turbo/5.1)
        # Ollama local descartado — VPS sem RAM suficiente (1.2GB < 1.7GB min)
        try:
            r = self._session.post(
                "https://api.z.ai/api/paas/v4/chat/completions",
                headers={"Authorization": f"Bearer {zai_key}", "Content-Type": "application/json"},
                json={
                    "model": "glm-5",
                    "messages": [{"role": "user", "content": message}],
                    "max_tokens": 1500,
                    "temperature": 0.3,
                },
                timeout=90,
            )
            r.raise_for_status()
            content = r.json()["choices"][0]["message"]["content"]
            logger.info("oxe_bridge: fallback-1 GLM/glm-5 -> %d chars", len(content))
            return content
        except Exception as exc:
            logger.warning("oxe_bridge: fallback-1 glm-5 falhou: %s", exc)

        # Fallback 2 + 3: Groq (provedor independente, < 1s latência)
        # Modelos testados e confirmados disponíveis: llama-3.3-70b, gpt-oss-120b
        groq_key = _os.environ.get("GROQ_API_KEY", "")
        if groq_key:
            for groq_model in ("llama-3.3-70b-versatile", "openai/gpt-oss-120b"):
                try:
                    r = self._session.post(
                        "https://api.groq.com/openai/v1/chat/completions",
                        headers={"Authorization": f"Bearer {groq_key}", "Content-Type": "application/json"},
                        json={
                            "model": groq_model,
                            "messages": [{"role": "user", "content": message}],
                            "max_tokens": 1500,
                            "temperature": 0.3,
                        },
                        timeout=30,
                    )
                    if r.status_code == 429:
                        logger.warning("oxe_bridge: fallback Groq/%s 429 — tentando próximo", groq_model)
                        continue
                    r.raise_for_status()
                    content = r.json()["choices"][0]["message"]["content"]
                    # Qwen3 retorna <think>...</think> — remover antes de retornar
                    import re as _re
                    content = _re.sub(r"<think>[\s\S]*?</think>", "", content).strip()
                    logger.info("oxe_bridge: fallback Groq/%s -> %d chars", groq_model, len(content))
                    return content
                except Exception as exc:
                    logger.warning("oxe_bridge: fallback Groq/%s falhou: %s", groq_model, exc)

        # Fallback 4: OpenRouter hermes-3-405b (último recurso, ~7s, provedor independente)
        or_key = _os.environ.get("OPENROUTER_API_KEY", "")
        if or_key:
            try:
                r = self._session.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers={"Authorization": f"Bearer {or_key}", "Content-Type": "application/json"},
                    json={
                        "model": "nousresearch/hermes-3-llama-3.1-405b:free",
                        "messages": [{"role": "user", "content": message}],
                        "max_tokens": 1500,
                        "temperature": 0.3,
                    },
                    timeout=45,
                )
                r.raise_for_status()
                content = r.json()["choices"][0]["message"]["content"]
                logger.info("oxe_bridge: fallback OpenRouter/hermes-405b -> %d chars", len(content))
                return content
            except Exception as exc:
                logger.warning("oxe_bridge: fallback OpenRouter falhou: %s", exc)

        return ""

    def escritorio_docx(self, conversation_id: str) -> bytes | None:
        """Baixa o .docx gerado pelo Escritório Virtual."""
        try:
            token = self._get_token()
            r = self._session.get(
                f"{self.oxe_url}/escritorio/{conversation_id}/docx",
                params={"authorization": f"Bearer {token}"},
                timeout=60,
            )
            r.raise_for_status()
            return r.content
        except Exception as exc:
            logger.warning("oxe_bridge: docx falhou: %s", exc)
            return None

    def render_docx_local(self, texto: str, titulo: str = "Peça Jurídica") -> bytes | None:
        """Gera .docx localmente a partir do texto — sem depender do Escritório Virtual."""
        try:
            import docx as _docx
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io

            doc = _docx.Document()
            for section in doc.sections:
                section.top_margin    = Cm(3)
                section.bottom_margin = Cm(2)
                section.left_margin   = Cm(3)
                section.right_margin  = Cm(2)

            style = doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)

            hdr = doc.sections[0].header
            hdr_p = hdr.paragraphs[0]
            hdr_p.text = "OXÉ Legal AI — Peça Gerada Automaticamente"
            hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if hdr_p.runs:
                hdr_p.runs[0].font.size = Pt(10)
                hdr_p.runs[0].bold = True

            t = doc.add_paragraph(titulo)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if t.runs:
                t.runs[0].bold = True
                t.runs[0].font.size = Pt(14)
            doc.add_paragraph()
            for line in texto.split(chr(10)):
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as exc:
            logger.warning("oxe_bridge: render_docx_local falhou: %s", exc)
            return None

    def health(self) -> dict:
        """Verifica saúde do bridge via estado interno (evita self-request circular)."""
        # Nota: NÃO fazemos HTTP request ao próprio servidor aqui.
        # O bridge roda no mesmo processo que a API — qualquer request HTTP
        # ao host bloquearia o event loop do uvicorn (single-worker).
        if not self._email or not self._password:
            return {"api": "offline", "reason": "credenciais ausentes"}
        if self.oxe_url:
            return {"api": "ok", "url": self.oxe_url}
        return {"api": "offline"}

    # ── Pipeline Premium ──────────────────────────────────────────────────────

    async def run_premium(
        self,
        query: str,
        conv_id: str | None = None,
        estilo_advogado: str = "",
        use_assembly: bool = False,
        adversarial_persona: str = "Procurador da Fazenda Nacional",
    ) -> PremiumResult:
        """
        Executa o pipeline premium completo (3 waves).

        Wave 1 — Inteligência (paralela, Ollama grátis):
          Busca semântica em jurisprudência, legislação e doutrina.

        Wave 2 — Estratégia & Construção (GLM/Claude):
          Jurimetria → Neuropersuasão → LexiaForge → JurisGuard.

        Wave 3 — Entrega (Ollama grátis):
          Formatação ABNT + .docx.
        """
        t0 = time.time()
        conv_id = conv_id or f"premium-{int(t0)}"
        result  = PremiumResult(query=query)
        custo   = 0.0

        logger.info("oxe_bridge: pipeline premium — '%s'", query[:80])

        # ── Wave 1: Inteligência (paralela) ──────────────────────────────────
        logger.info("oxe_bridge: Wave 1 — pesquisa paralela")

        loop = asyncio.get_running_loop()
        juris_f = loop.run_in_executor(None, self.buscar, query, "oxe_jurisprudencia", 5)
        legis_f = loop.run_in_executor(None, self.buscar, query, "oxe_legislacao", 5)
        doutr_f = loop.run_in_executor(None, self.buscar, query, "oxe_regulatorio", 3)

        (
            result.wave1_jurisprudencia,
            result.wave1_legislacao,
            result.wave1_doutrina,
        ) = await asyncio.gather(juris_f, legis_f, doutr_f)

        logger.info(
            "oxe_bridge: Wave 1 concluída — juris=%d, legis=%d, doutr=%d",
            len(result.wave1_jurisprudencia),
            len(result.wave1_legislacao),
            len(result.wave1_doutrina),
        )

        # ── Wave 2: Estratégia & Construção ──────────────────────────────────
        # Dois modos: sequential (padrão) ou assembly (LegalAssembly paralelo)

        if use_assembly:
            # ── Wave 2 ASSEMBLY MODE — LegalAssembly paralelo ─────────────────
            logger.info("oxe_bridge: Wave 2 — LegalAssembly (modo paralelo)")
            result.wave2_mode = "assembly"
            try:
                from symbiont.legal_assembly import LegalAssembly
                from symbiont.translation import TranslationLayer
                import os as _os

                _translation = TranslationLayer(
                    groq_key=_os.environ.get("GROQ_API_KEY", ""),
                    zai_key=_os.environ.get("ZAI_API_KEY", ""),
                )
                _assembly = LegalAssembly(
                    groq_key=_os.environ.get("GROQ_API_KEY", ""),
                    zai_key=_os.environ.get("ZAI_API_KEY", ""),
                    openrouter_key=_os.environ.get("OPENROUTER_API_KEY", ""),
                    translation_layer=_translation,
                )
                _shared_ctx = {
                    "jurisprudencia": result.wave1_jurisprudencia,
                    "legislacao":     result.wave1_legislacao,
                    "doutrina":       result.wave1_doutrina,
                }
                _asm_result = await _assembly.run(
                    query=query,
                    shared_context=_shared_ctx,
                    adversarial_persona=adversarial_persona,
                    roles=["jurimetria", "argumentacao", "precedentes"],
                    conv_id=conv_id,
                )
                # Mapear campos assembly → campos PremiumResult
                result.wave2_jurimetria           = _asm_result.jurimetria
                result.wave2_persuasao            = _asm_result.argumentacao
                result.wave2_rascunho             = _asm_result.sintese
                result.wave2_revisado             = _asm_result.refinado
                result.wave2_assembly_sintese     = _asm_result.sintese
                result.wave2_assembly_adversarial = _asm_result.adversarial
                result.wave2_assembly_quality     = _asm_result.quality_score
                result.wave2_assembly_models      = _asm_result.models_used
                if _asm_result.erros:
                    result.erros.extend([f"assembly:{e}" for e in _asm_result.erros])
                custo += 0.040   # estimativa: 6 chamadas Groq/ZAI
                logger.info(
                    "oxe_bridge: Wave 2 assembly concluída — quality=%.2f, models=%s, erros=%d",
                    _asm_result.quality_score,
                    list(_asm_result.models_used.values()),
                    len(_asm_result.erros),
                )
            except Exception as _e:
                logger.error("oxe_bridge: Wave 2 assembly falhou (%s) — fallback sequential", _e)
                result.erros.append(f"assembly_fallback: {_e}")
                result.wave2_mode = "sequential_fallback"
                use_assembly = False   # cair no bloco sequential abaixo

        if not use_assembly:
            # ── Wave 2 SEQUENTIAL MODE — pipeline original 2A→2B→2C→2D ────────
            logger.info("oxe_bridge: Wave 2 — modo sequencial")
            result.wave2_mode = result.wave2_mode or "sequential"

            # Wave 2A: Jurimetria
            logger.info("oxe_bridge: Wave 2A — Jurimetria")
            juris_context = "\n".join(
                f"- {r.get('tribunal','?')} ({r.get('data','?')}): {r.get('ementa','')[:80]}"
                for r in result.wave1_jurisprudencia
            )
            juris_prompt = (
                f"QUERY: {query}\n\n"
                f"JURISPRUDÊNCIA ENCONTRADA:\n{juris_context}\n\n"
                "Analise a probabilidade de sucesso desta tese considerando:\n"
                "1. Tendência predominante dos tribunais\n"
                "2. Análise por câmara/turma quando possível\n"
                "3. Pontos fortes e fracos da posição\n"
                "Seja objetivo e quantitativo quando possível."
            )
            result.wave2_jurimetria = await loop.run_in_executor(
                None, self.chat, juris_prompt, None, f"{conv_id}-juri", "medium"
            )
            custo += COST_TIERS["cheap"]["cost_usd"]

            # Wave 2B: Neuropersuasão
            logger.info("oxe_bridge: Wave 2B — Neuropersuasão")
            persuasao_prompt = (
                f"QUERY: {query}\n\n"
                f"ANÁLISE JURIMETRIA:\n{result.wave2_jurimetria[:300]}\n\n"
                "Identifique:\n"
                "1. Tom argumentativo ideal para este juízo (técnico/emocional/assertivo)\n"
                "2. Elementos retóricos mais eficazes baseados nos julgados\n"
                "3. Argumentos que mais ressoam com este tribunal\n"
                "4. Gatilhos cognitivos que devem ser ativados\n"
                "Baseie-se nos padrões reais das decisões encontradas."
            )
            result.wave2_persuasao = await loop.run_in_executor(
                None, self.chat, persuasao_prompt, None, f"{conv_id}-persuasao", "medium"
            )
            custo += COST_TIERS["medium"]["cost_usd"]

            # Wave 2C: Redação (LexiaForge)
            logger.info("oxe_bridge: Wave 2C — LexiaForge redação")
            legis_context = "\n".join(
                f"- {r.get('titulo','?')}: {r.get('ementa','')[:60]}"
                for r in result.wave1_legislacao
            )
            draft_prompt = (
                f"DEMANDA: {query}\n\n"
                f"JURISPRUDÊNCIA:\n{juris_context}\n\n"
                f"LEGISLAÇÃO APLICÁVEL:\n{legis_context}\n\n"
                f"ANÁLISE JURIMETRIA:\n{result.wave2_jurimetria[:250]}\n\n"
                f"CALIBRAÇÃO RETÓRICA:\n{result.wave2_persuasao[:200]}\n\n"
                + (f"ESTILO DO ADVOGADO:\n{estilo_advogado}\n\n" if estilo_advogado else "")
                + "Redija a peça jurídica completa em prosa fluida, "
                  "incorporando todos os elementos acima. "
                  "Estrutura: Preâmbulo → Fatos → Fundamentos Jurídicos → Pedidos."
            )
            result.wave2_rascunho = await loop.run_in_executor(
                None, self.chat, draft_prompt, "lexiaforge", f"{conv_id}-draft", "high"
            )
            custo += COST_TIERS["high"]["cost_usd"]

            # Wave 2D: Revisão (JurisGuard)
            logger.info("oxe_bridge: Wave 2D — JurisGuard revisão")
            review_prompt = (
                "Revise e fortaleça a peça abaixo, atuando como advogado adverso "
                "que tentará atacá-la. Identifique e corrija todos os pontos fracos.\n\n"
                f"PEÇA ORIGINAL:\n{result.wave2_rascunho[:2000]}"
            )
            result.wave2_revisado = await loop.run_in_executor(
                None, self.chat, review_prompt, None, f"{conv_id}-review", "medium"
            )
            custo += COST_TIERS["cheap"]["cost_usd"]

        # ── Wave 3: Formatação & Entrega ──────────────────────────────────────
        logger.info("oxe_bridge: Wave 3 — formatação")

        await loop.run_in_executor(
            None,
            self.chat,
            (
                "Formate e estruture a seguinte peça jurídica conforme padrão ABNT "
                f"e templates de grandes escritórios de SP:\n\n{result.wave2_revisado[:2500]}"
            ),
            None,
            f"{conv_id}-format",
            "medium",
        )

        docx_bytes = await loop.run_in_executor(
            None, self.render_docx_local, result.wave3_peca, "Peça Jurídica"
        )
        if docx_bytes:
            docx_path = os.path.join(
                tempfile.gettempdir(), f"oxe_premium_{conv_id}.docx"
            )
            with open(docx_path, "wb") as f:
                f.write(docx_bytes)
            result.docx_path = docx_path
            logger.info("oxe_bridge: .docx salvo em %s", docx_path)

        result.wave3_peca         = result.wave2_revisado
        result.custo_estimado_usd = custo
        result.tempo_total_s      = time.time() - t0

        logger.info(
            "oxe_bridge: pipeline concluído — tempo=%.1fs, custo=$%.3f, docx=%s",
            result.tempo_total_s,
            result.custo_estimado_usd,
            "✓" if result.docx_path else "✗",
        )
        return result


# ─── Router FastAPI (opcional) ────────────────────────────────────────────────

def create_premium_router(bridge: OXEBridge | None = None):
    """
    Cria o router FastAPI /premium para integração no gateway OXÉ.

    Uso em chat.py:
        from symbiont.oxe_bridge import create_premium_router
        app.include_router(create_premium_router())
    """
    try:
        from fastapi import APIRouter, Body, HTTPException
    except ImportError:
        logger.warning("oxe_bridge: fastapi não instalado — router não disponível")
        return None

    router  = APIRouter(prefix="/premium", tags=["premium"])
    _bridge = bridge or OXEBridge()

    @router.post("")
    async def premium_endpoint(req: PremiumRequest = Body(...)):
        """Pipeline premium multiagente (3 waves). Retorna peça completa + path .docx."""
        if not req.query.strip():
            raise HTTPException(400, "query obrigatória")
        try:
            result = await _bridge.run_premium(
                query=req.query,
                conv_id=req.conversation_id or None,
                estilo_advogado=req.estilo_advogado,
                use_assembly=req.use_assembly,
                adversarial_persona=req.adversarial_persona,
            )
            return result.to_dict()
        except Exception as exc:
            raise HTTPException(500, str(exc))

    @router.get("/health")
    async def premium_health():
        return {"bridge": "ok", "oxe": _bridge.health()}

    return router
